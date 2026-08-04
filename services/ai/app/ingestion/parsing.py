"""PDF and DOCX parsing.

Produces blocks carrying enough position to make a citation useful: which page
the text came from, and which headings it sat under. A citation that can only
say "somewhere in this document" is not worth showing.

Failures raise `ParseError` with a stable code. Parser exception text can quote
the document it failed on, and these codes are stored and shown to users, so the
underlying message never travels with them.
"""

from __future__ import annotations

import io
from collections import Counter
from dataclasses import dataclass, field
from typing import Literal

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

ParseFailureCode = Literal[
    "FILE_UNREADABLE",
    "FILE_ENCRYPTED",
    "FILE_EMPTY",
    "UNSUPPORTED_MEDIA_TYPE",
]

PDF_MEDIA_TYPE = "application/pdf"
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# A line whose largest glyph is at least this much bigger than the document's
# usual body text is treated as a heading. Tuned to catch real section titles
# without promoting mildly emphasised words.
HEADING_SIZE_RATIO = 1.15
# Headings are short. Without this, a whole paragraph set in a slightly larger
# face becomes a heading and swallows the section structure.
HEADING_MAX_CHARS = 120


class ParseError(Exception):
    """A document that cannot be turned into text, with a safe reason."""

    def __init__(self, code: ParseFailureCode, message: str) -> None:
        super().__init__(message)
        self.code: ParseFailureCode = code


@dataclass(frozen=True)
class ParsedBlock:
    """One paragraph-sized run of text, with where it came from."""

    ordinal: int
    text: str
    page_number: int | None
    heading_path: tuple[str, ...] = field(default=())


@dataclass(frozen=True)
class ParseResult:
    blocks: tuple[ParsedBlock, ...]
    page_count: int


def _normalise(text: str) -> str:
    """Collapse whitespace without joining separate words."""
    return " ".join(text.split())


def _parse_pdf(data: bytes) -> ParseResult:
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as error:
        raise ParseError("FILE_UNREADABLE", "The PDF could not be opened.") from error

    with document:
        if document.needs_pass:
            raise ParseError("FILE_ENCRYPTED", "The PDF is password protected.")

        # Two passes. The first learns what body text looks like in *this*
        # document, because absolute font sizes say nothing on their own — a
        # 14pt line is a heading in one document and body text in another.
        # Size -> how many characters are set at that size. Weighting by
        # characters rather than by line count is what makes this robust: body
        # text always dominates a document by volume, even in one with many
        # short headings, whereas a plain median is dragged upward as soon as
        # headings are not heavily outnumbered.
        size_weights: Counter[float] = Counter()
        pages: list[list[tuple[str, float]]] = []

        for page in document:
            lines: list[tuple[str, float]] = []
            blocks = page.get_text("dict").get("blocks", [])

            for block in blocks:
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    text = _normalise("".join(span.get("text", "") for span in spans))

                    if not text:
                        continue

                    largest = max((float(span.get("size", 0.0)) for span in spans), default=0.0)
                    lines.append((text, largest))
                    # Round so that sizes differing by hinting noise group.
                    size_weights[round(largest, 1)] += len(text)

            pages.append(lines)

        if not size_weights:
            raise ParseError("FILE_EMPTY", "The PDF contains no extractable text.")

        body_size = size_weights.most_common(1)[0][0]
        heading_threshold = body_size * HEADING_SIZE_RATIO

        blocks_out: list[ParsedBlock] = []
        current_heading: tuple[str, ...] = ()
        ordinal = 0

        for page_index, lines in enumerate(pages, start=1):
            for text, size in lines:
                if size >= heading_threshold and len(text) <= HEADING_MAX_CHARS:
                    # PDFs carry no heading levels, only visual weight, so the
                    # path stays one deep rather than inventing a hierarchy.
                    current_heading = (text,)
                    continue

                ordinal += 1
                blocks_out.append(
                    ParsedBlock(
                        ordinal=ordinal,
                        text=text,
                        page_number=page_index,
                        heading_path=current_heading,
                    )
                )

        if not blocks_out:
            raise ParseError("FILE_EMPTY", "The PDF contains no body text.")

        return ParseResult(blocks=tuple(blocks_out), page_count=len(pages))


def _heading_level(style_name: str) -> int | None:
    """The depth of a Word heading style, or None if it is not a heading."""
    if not style_name.startswith("Heading"):
        return None

    suffix = style_name.removeprefix("Heading").strip()

    try:
        return int(suffix)
    except ValueError:
        # "Heading" with no number, or a localised style name.
        return 1


def _parse_docx(data: bytes) -> ParseResult:
    try:
        document = DocxDocument(io.BytesIO(data))
    except PackageNotFoundError as error:
        raise ParseError("FILE_UNREADABLE", "The DOCX could not be opened.") from error
    except Exception as error:
        raise ParseError("FILE_UNREADABLE", "The DOCX could not be read.") from error

    blocks_out: list[ParsedBlock] = []
    # Word gives real heading levels, so the path can be a genuine hierarchy
    # rather than the flat approximation the PDF path settles for.
    heading_stack: list[str] = []
    ordinal = 0

    for paragraph in document.paragraphs:
        text = _normalise(paragraph.text)

        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style is not None else ""
        level = _heading_level(style_name or "")

        if level is not None:
            del heading_stack[level - 1 :]
            heading_stack.append(text)
            continue

        ordinal += 1
        blocks_out.append(
            ParsedBlock(
                ordinal=ordinal,
                text=text,
                # DOCX has no page concept before rendering. Claiming a page
                # number here would put a number on a citation that is wrong.
                page_number=None,
                heading_path=tuple(heading_stack),
            )
        )

    # Tables often hold the onboarding detail people actually ask about, so
    # they are captured rather than dropped, flattened row by row.
    for table in document.tables:
        for row in table.rows:
            cells = [_normalise(cell.text) for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)

            if not text:
                continue

            ordinal += 1
            blocks_out.append(
                ParsedBlock(
                    ordinal=ordinal,
                    text=text,
                    page_number=None,
                    heading_path=tuple(heading_stack),
                )
            )

    if not blocks_out:
        raise ParseError("FILE_EMPTY", "The DOCX contains no text.")

    return ParseResult(blocks=tuple(blocks_out), page_count=0)


def parse_document(data: bytes, media_type: str) -> ParseResult:
    """Parse bytes of a supported document into positioned blocks."""
    if not data:
        raise ParseError("FILE_EMPTY", "The file is empty.")

    if media_type == PDF_MEDIA_TYPE:
        return _parse_pdf(data)

    if media_type == DOCX_MEDIA_TYPE:
        return _parse_docx(data)

    raise ParseError("UNSUPPORTED_MEDIA_TYPE", "Only PDF and DOCX are supported.")
