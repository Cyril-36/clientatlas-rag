"""Parsing, chunking and embedding.

Documents are generated rather than committed as fixtures: a checked-in binary
is opaque in review, and nobody notices when it stops representing the case it
was added for.
"""

from __future__ import annotations

import base64
import io

import fitz
import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from app.embedding.provider import DeterministicProvider
from app.ingestion.chunking import EMBEDDING_WINDOW_TOKENS, chunk_blocks, estimate_tokens
from app.ingestion.parsing import (
    DOCX_MEDIA_TYPE,
    PDF_MEDIA_TYPE,
    ParsedBlock,
    ParseError,
    parse_document,
)
from app.main import create_app

client = TestClient(create_app())


def build_pdf(sections: list[tuple[str, list[str]]]) -> bytes:
    """A PDF with visually larger headings, as a real document would have."""
    document = fitz.open()
    page = document.new_page()
    y = 72.0

    for heading, paragraphs in sections:
        page.insert_text((72, y), heading, fontsize=18)
        y += 30

        for paragraph in paragraphs:
            page.insert_text((72, y), paragraph, fontsize=11)
            y += 18

            if y > 720:
                page = document.new_page()
                y = 72.0

    data: bytes = document.tobytes()
    document.close()
    return data


def build_docx(sections: list[tuple[str, list[str]]]) -> bytes:
    document = DocxDocument()

    for heading, paragraphs in sections:
        document.add_heading(heading, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestPdfParsing:
    def test_extracts_text_with_page_numbers(self) -> None:
        pdf = build_pdf([("Security", ["Access requests go through the platform team."])])

        result = parse_document(pdf, PDF_MEDIA_TYPE)

        assert result.page_count == 1
        assert any("platform team" in block.text for block in result.blocks)
        assert all(block.page_number == 1 for block in result.blocks)

    def test_recognises_headings_by_relative_size(self) -> None:
        # The heading is larger than body text in this document; an absolute
        # threshold would be wrong for a document set in a different size.
        pdf = build_pdf([("Access requirements", ["Ask your manager first."])])

        result = parse_document(pdf, PDF_MEDIA_TYPE)
        body = [block for block in result.blocks if "manager" in block.text]

        assert body
        assert body[0].heading_path == ("Access requirements",)

    def test_rejects_bytes_that_are_not_a_pdf(self) -> None:
        with pytest.raises(ParseError) as caught:
            parse_document(b"not a pdf at all", PDF_MEDIA_TYPE)

        assert caught.value.code == "FILE_UNREADABLE"

    def test_rejects_a_pdf_with_no_text(self) -> None:
        empty = fitz.open()
        empty.new_page()
        data: bytes = empty.tobytes()
        empty.close()

        with pytest.raises(ParseError) as caught:
            parse_document(data, PDF_MEDIA_TYPE)

        assert caught.value.code == "FILE_EMPTY"


class TestDocxParsing:
    def test_builds_a_heading_hierarchy(self) -> None:
        docx = build_docx([("Onboarding", ["Day one is orientation."])])

        result = parse_document(docx, DOCX_MEDIA_TYPE)
        body = [block for block in result.blocks if "orientation" in block.text]

        assert body
        assert body[0].heading_path == ("Onboarding",)

    def test_reports_no_page_number(self) -> None:
        # DOCX has no pages before rendering. A number here would be a citation
        # pointing somewhere the reader cannot find.
        docx = build_docx([("Onboarding", ["Day one is orientation."])])

        result = parse_document(docx, DOCX_MEDIA_TYPE)

        assert all(block.page_number is None for block in result.blocks)

    def test_captures_table_rows(self) -> None:
        document = DocxDocument()
        document.add_heading("Contacts", level=1)
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Platform team"
        table.rows[0].cells[1].text = "platform@example.test"
        buffer = io.BytesIO()
        document.save(buffer)

        result = parse_document(buffer.getvalue(), DOCX_MEDIA_TYPE)

        assert any("platform@example.test" in block.text for block in result.blocks)

    def test_rejects_bytes_that_are_not_a_docx(self) -> None:
        with pytest.raises(ParseError) as caught:
            parse_document(b"PK\x03\x04 but not really", DOCX_MEDIA_TYPE)

        assert caught.value.code == "FILE_UNREADABLE"


class TestChunking:
    def block(self, ordinal: int, text: str, heading: str = "Section") -> ParsedBlock:
        return ParsedBlock(ordinal=ordinal, text=text, page_number=1, heading_path=(heading,))

    def test_groups_blocks_up_to_the_target(self) -> None:
        blocks = [self.block(index, "word " * 50) for index in range(1, 21)]

        chunks = chunk_blocks(blocks, target_tokens=200, max_tokens=250, overlap_tokens=30)

        assert len(chunks) > 1
        assert all(chunk.token_count <= 300 for chunk in chunks)

    def test_never_spans_a_heading_boundary(self) -> None:
        blocks = [
            self.block(1, "Alpha content here.", heading="Alpha"),
            self.block(2, "Beta content here.", heading="Beta"),
        ]

        chunks = chunk_blocks(blocks, target_tokens=500, max_tokens=800)

        # Both are tiny and would otherwise merge, but mixing sections retrieves
        # for queries about either and answers neither well.
        assert len(chunks) == 2
        assert chunks[0].heading_path == ("Alpha",)
        assert chunks[1].heading_path == ("Beta",)

    def test_overlaps_consecutive_chunks_within_a_section(self) -> None:
        blocks = [self.block(index, f"sentence number {index} " * 20) for index in range(1, 16)]

        chunks = chunk_blocks(blocks, target_tokens=120, max_tokens=160, overlap_tokens=60)

        assert len(chunks) >= 2
        # A fact at the end of one chunk stays findable from the next.
        tail = chunks[0].text.split("\n")[-1]
        assert tail in chunks[1].text

    def test_does_not_overlap_across_a_heading_boundary(self) -> None:
        blocks = [
            self.block(1, "alpha " * 100, heading="Alpha"),
            self.block(2, "beta " * 100, heading="Beta"),
        ]

        chunks = chunk_blocks(blocks, target_tokens=50, max_tokens=80, overlap_tokens=40)

        for chunk in chunks:
            if chunk.heading_path == ("Beta",):
                assert "alpha" not in chunk.text

    def test_terminates_on_a_single_oversized_block(self) -> None:
        # A block larger than max_tokens cannot be split by this chunker, and a
        # naive overlap would carry it forward for ever.
        chunks = chunk_blocks([self.block(1, "word " * 5000)], target_tokens=100, max_tokens=150)

        assert len(chunks) == 1

    def test_ordinals_are_contiguous_from_one(self) -> None:
        blocks = [self.block(index, "word " * 60) for index in range(1, 11)]

        chunks = chunk_blocks(blocks, target_tokens=100, max_tokens=150, overlap_tokens=20)

        assert [chunk.ordinal for chunk in chunks] == list(range(1, len(chunks) + 1))

    def test_estimate_tokens_is_never_zero(self) -> None:
        assert estimate_tokens("") >= 1
        assert estimate_tokens("one") >= 1


class TestDeterministicProvider:
    def test_produces_correctly_shaped_vectors(self) -> None:
        provider = DeterministicProvider()

        vectors = provider.embed(["hello", "world"])

        assert len(vectors) == 2
        assert all(len(vector) == 384 for vector in vectors)

    def test_is_stable_for_the_same_input(self) -> None:
        provider = DeterministicProvider()

        assert provider.embed(["same"]) == provider.embed(["same"])

    def test_differs_for_different_input(self) -> None:
        provider = DeterministicProvider()

        assert provider.embed(["one"]) != provider.embed(["two"])

    def test_is_normalised(self) -> None:
        provider = DeterministicProvider()

        [vector] = provider.embed(["anything"])
        magnitude = sum(value * value for value in vector) ** 0.5

        assert magnitude == pytest.approx(1.0, abs=1e-6)

    def test_does_not_claim_to_be_the_real_model(self) -> None:
        # The Node contract pins the model name, so a stored vector can never be
        # one of these by accident.
        assert DeterministicProvider().model_name != "sentence-transformers/all-MiniLM-L6-v2"


class TestIngestionEndpoints:
    def test_parse_returns_blocks_and_chunks(self) -> None:
        pdf = build_pdf([("Security", ["Access requests go through the platform team."])])

        response = client.post(
            "/v1/parse",
            json={
                "filename": "handbook.pdf",
                "mediaType": PDF_MEDIA_TYPE,
                "contentBase64": base64.b64encode(pdf).decode(),
            },
        )

        assert response.status_code == 200
        body = response.json()
        assert body["pageCount"] == 1
        assert body["blocks"]
        assert body["chunks"]
        assert len(body["checksumSha256"]) == 64

    def test_parse_reports_a_code_without_parser_detail(self) -> None:
        response = client.post(
            "/v1/parse",
            json={
                "filename": "broken.pdf",
                "mediaType": PDF_MEDIA_TYPE,
                "contentBase64": base64.b64encode(b"nonsense").decode(),
            },
        )

        assert response.status_code == 422
        assert response.json()["detail"] == {"code": "FILE_UNREADABLE"}

    def test_parse_rejects_invalid_base64(self) -> None:
        response = client.post(
            "/v1/parse",
            json={
                "filename": "x.pdf",
                "mediaType": PDF_MEDIA_TYPE,
                "contentBase64": "!!!! not base64 !!!!",
            },
        )

        assert response.status_code == 422

    def test_embed_reports_the_provider_that_actually_ran(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.setenv("CLIENTATLAS_AI_EMBEDDING_PROVIDER", "deterministic")
        from app.config import get_settings

        get_settings.cache_clear()

        response = client.post("/v1/embed", json={"texts": ["hello"]})

        assert response.status_code == 200
        body = response.json()
        assert body["dimensions"] == 384
        assert len(body["vectors"][0]) == 384
        # Not the real model name — the caller's contract will refuse it.
        assert body["model"] == "deterministic-test-provider"

        get_settings.cache_clear()

    def test_embed_rejects_an_oversized_batch(self) -> None:
        response = client.post("/v1/embed", json={"texts": ["x"] * 65})

        assert response.status_code == 422


class TestEmbeddingWindowExposure:
    """How much chunk text the encoder never sees.

    all-MiniLM-L6-v2 reads the first 256 WordPiece tokens and silently discards
    the rest, so a chunk longer than that is partly invisible to retrieval with
    no error anywhere. EMBEDDING_WINDOW_TOKENS existed as documentation only;
    this makes it load-bearing.

    These record the current exposure rather than forbidding it. Chunks are cut
    at heading boundaries, not at the encoder window, and a measured sweep showed
    that resizing chunks to fit the window makes retrieval worse — so the
    exposure is accepted, and this is what stops it growing unnoticed.
    """

    def test_the_window_constant_matches_the_model(self) -> None:
        # If the model is ever swapped, this is the line that should fail first.
        assert EMBEDDING_WINDOW_TOKENS == 256

    def test_most_chunks_fit_inside_the_window(self) -> None:
        blocks = [
            ParsedBlock(
                ordinal=i,
                text=f"Sentence {i} about access policy. " * 4,
                page_number=1,
                heading_path=(f"Section {i // 5}",),
            )
            for i in range(1, 60)
        ]

        chunks = chunk_blocks(blocks)
        over = [c for c in chunks if c.token_count > EMBEDDING_WINDOW_TOKENS]

        # A majority fitting is the property that matters; the tail is known.
        assert len(over) / len(chunks) < 0.5

    def test_an_oversized_paragraph_is_not_split_and_is_known_to_truncate(self) -> None:
        # A single block longer than the window cannot be split by a chunker that
        # only groups whole blocks. Recorded here so the limitation is visible,
        # rather than discovered later as a retrieval mystery.
        huge = ParsedBlock(ordinal=1, text="word " * 900, page_number=1, heading_path=("H",))

        [chunk] = chunk_blocks([huge])

        assert chunk.token_count > EMBEDDING_WINDOW_TOKENS
